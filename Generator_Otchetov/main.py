from docx import Document
from docx.shared import Inches

document = Document()

NomerRabot = input("Nomer work: ")
NomerGroop = input("Nomer group: ")
Prepod = input("Teacher: ")
Student = input("Student: ")
NameWork = input("Name work: ")
Goal = input("Goal: ")
Code = input("Code: ")
PostanovkaZadachi = input("Zadacha: ")
Test = input("Test: ")
End = input("End: ")
Information = input("Information: ")

with open(Code, "r") as file:
    for line in file:
        Text = line

paragraph = document.add_paragraph('Номер отчета: '+NomerRabot)
paragraph = document.add_paragraph('Номер группы: '+NomerGroop)
paragraph = document.add_paragraph('Преподаватель: '+Prepod)
paragraph = document.add_paragraph('Студент: '+Student)
paragraph = document.add_paragraph('Название работы: '+NameWork)
paragraph = document.add_paragraph('Цель работы: '+Goal)
paragraph = document.add_paragraph('Код: '+Text)
paragraph = document.add_paragraph('Тест: ')
document.add_picture(Test, width=Inches(1.0))
paragraph = document.add_paragraph('Вывод: '+End)
paragraph = document.add_paragraph('Источники: '+Information)

document.save('test1.docx')